"""
📄 Word 文档生成 Agent（支持用户自定义模板）

功能：
- 根据用户输入的主题/内容，自动生成格式精美的 Word 文档
- 支持用户上传自定义 .docx 模板，保留模板样式和布局
- 模板占位符替换：模板中的 {{变量名}} 自动替换为用户提供的值
- 支持自定义标题、章节、段落、列表、表格
- 支持设置文档风格（字体、字号、颜色、对齐方式）
- 支持添加页眉页脚、页码
- 输出可下载的 .docx 文件

架构说明：
- 使用 agentic 模式，LLM 节点通过工具调用动态生成文档
- 模板发现 → 内容规划 → 文档生成，三步证据流
- 工具层面操作文件系统，验证后写入

使用模板的两种方式：
1. 传入 template_path 参数指向已有 .docx 文件
2. 不传模板时自动使用内置样式生成

Run:
    agentclaw up  # 启动服务后可通过 API 调用
"""

import os
import re
import json
import shutil
import asyncio
from datetime import datetime
from pathlib import Path
from agentclaw import Workflow, LLMNode, Input

# ============================================================
# Workflow 定义
# ============================================================
workflow = Workflow(
    id="docx_generator",
    name="📄 Word 文档生成器（支持模板）",
    description="根据用户输入的主题和内容，自动生成格式精美的 Word 文档（.docx），支持用户上传自定义模板，保留模板样式和布局。",
    version="2.0.0",
    welcome="📄 你好！我是 Word 文档生成助手！你可以提供自己的 .docx 模板文件，我会保留你的样式并填充内容；不传模板我也能用内置样式帮你生成精美文档！",
    timeout=240,
    inputs=[
        Input("topic", str, required=True,
              description="文档主题，例如：'项目周报'、'产品调研报告'、'会议纪要'"),
        Input("content", str, required=False, default="",
              description="文档详细内容或大纲，可以包含标题、段落、要点等（可选，不填则由 AI 自动生成）"),
        Input("template_path", str, required=False, default="",
              description="自定义模板文件路径（.docx），模板中的 {{变量名}} 占位符会被替换（可选）"),
        Input("template_values", str, required=False, default="",
              description="模板占位符替换值，JSON 格式，如：{\"公司名\": \"ABC科技\", \"项目名\": \"X计划\"}（可选）"),
        Input("author", str, required=False, default="",
              description="文档作者名称（可选）"),
        Input("style", str, required=False, default="professional",
              description="文档风格（不使用模板时生效）：professional（专业）/ creative（创意）/ simple（简洁）",
              choices=["professional", "creative", "simple"]),
        Input("include_toc", str, required=False, default="yes",
              description="是否包含目录（不使用模板时生效）：yes / no",
              choices=["yes", "no"]),
        Input("include_page_numbers", str, required=False, default="yes",
              description="是否包含页码（不使用模板时生效）：yes / no",
              choices=["yes", "no"]),
        Input("filename", str, required=False, default="",
              description="输出文件名（不含扩展名，可选，默认自动生成）"),
    ],
    user_input="topic",
)

# ============================================================
# 目录配置
# ============================================================
OUTPUT_DIR = Path("generated_docs")
TEMPLATES_DIR = Path("user_templates")
OUTPUT_DIR.mkdir(exist_ok=True)
TEMPLATES_DIR.mkdir(exist_ok=True)

# ============================================================
# 工具 1: 分析模板文件
# ============================================================
@workflow.tool
def analyze_template(template_path: str = "", **kwargs):
    """分析 .docx 模板文件，提取样式信息、占位符和结构。

    Args:
        template_path: 模板文件路径（相对于项目目录或绝对路径）
    Returns:
        模板分析结果：样式信息、占位符列表、段落结构概览
    """
    from docx import Document

    # 解析路径
    tpl_path = Path(template_path)
    if not tpl_path.is_absolute():
        # 尝试在项目目录、user_templates 目录、generated_docs 目录查找
        candidates = [
            Path.cwd() / template_path,
            TEMPLATES_DIR / template_path,
            TEMPLATES_DIR / f"{template_path}.docx",
            OUTPUT_DIR / template_path,
        ]
        for c in candidates:
            if c.exists():
                tpl_path = c
                break

    if not tpl_path.exists():
        return {"error": f"模板文件不存在: {tpl_path}", "found": False}

    doc = Document(str(tpl_path))

    # 提取占位符 {{变量名}}
    placeholders = set()
    placeholder_positions = []
    for i, para in enumerate(doc.paragraphs):
        matches = re.findall(r'\{\{(\w+)\}\}', para.text)
        for m in matches:
            placeholders.add(m)
            placeholder_positions.append({
                "placeholder": m,
                "paragraph_index": i,
                "text_snippet": para.text[:80]
            })

    # 提取段落样式概况
    style_summary = {}
    for para in doc.paragraphs:
        s = para.style.name if para.style else "Normal"
        style_summary[s] = style_summary.get(s, 0) + 1

    # 提取表格信息
    table_count = len(doc.tables)

    # 提取节/页面设置
    sections_info = []
    for sec in doc.sections:
        sections_info.append({
            "page_width": str(sec.page_width) if sec.page_width else "default",
            "page_height": str(sec.page_height) if sec.page_height else "default",
            "orientation": "landscape" if sec.orientation else "portrait",
        })

    # 段落数量
    para_count = len(doc.paragraphs)

    return {
        "found": True,
        "filename": tpl_path.name,
        "path": str(tpl_path),
        "paragraph_count": para_count,
        "table_count": table_count,
        "section_count": len(sections_info),
        "styles_used": style_summary,
        "placeholders_found": sorted(list(placeholders)),
        "placeholder_details": placeholder_positions[:20],  # 最多前20个
        "sections": sections_info,
        "has_content": para_count > 1,
    }


# ============================================================
# 工具 2: 基于模板生成文档（替换占位符 + 追加内容）
# ============================================================
@workflow.tool
def generate_from_template(
    template_path: str,
    topic: str,
    sections: list,
    template_values: dict = None,
    author: str = "",
    filename: str = "",
    **kwargs
):
    """基于 .docx 模板生成文档：保留模板样式，替换占位符，追加内容。

    Args:
        template_path: 模板文件路径
        topic: 文档标题/主题
        sections: 文档章节列表（追加到模板末尾），每个元素为 {"heading":..., "level":..., "type":..., "content":...}
        template_values: 模板占位符替换值字典，如 {"公司名": "ABC科技"}
        author: 作者名称
        filename: 输出文件名（不含扩展名）
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    # 解析模板路径
    tpl_path = Path(template_path)
    if not tpl_path.is_absolute():
        candidates = [
            Path.cwd() / template_path,
            TEMPLATES_DIR / template_path,
            TEMPLATES_DIR / f"{template_path}.docx",
            OUTPUT_DIR / template_path,
        ]
        for c in candidates:
            if c.exists():
                tpl_path = c
                break

    if not tpl_path.exists():
        return {"error": f"模板文件不存在: {tpl_path}"}

    # 基于模板创建文档
    doc = Document(str(tpl_path))

    # ---- 替换模板中的占位符 ----
    if template_values:
        for para in doc.paragraphs:
            for key, val in template_values.items():
                placeholder = "{{" + key + "}}"
                if placeholder in para.text:
                    # 保留原有格式替换文本
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(val))
                    # 如果 run 级别没替换完，整段替换
                    if placeholder in para.text:
                        para.text = para.text.replace(placeholder, str(val))

        # 表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for key, val in template_values.items():
                            placeholder = "{{" + key + "}}"
                            if placeholder in para.text:
                                for run in para.runs:
                                    if placeholder in run.text:
                                        run.text = run.text.replace(placeholder, str(val))
                                if placeholder in para.text:
                                    para.text = para.text.replace(placeholder, str(val))

    # ---- 获取模板的样式参考 ----
    # 尝试从模板中找标题样式
    template_styles = {}
    for para in doc.paragraphs:
        if para.style:
            sname = para.style.name
            if sname.startswith("Heading") or sname.startswith("heading"):
                level = sname[-1] if sname[-1].isdigit() else "1"
                template_styles[f"heading{level}"] = sname

    # ---- 辅助函数 ----
    def add_styled_paragraph(text, style_name=None, font_name=None, size=None,
                              bold=None, color=None, alignment=None,
                              space_before=0, space_after=6):
        p = doc.add_paragraph()
        if style_name:
            try:
                p.style = doc.styles[style_name]
            except:
                pass
        if alignment:
            p.alignment = alignment
        pf = p.paragraph_format
        pf.space_before = Pt(space_before)
        pf.space_after = Pt(space_after)
        run = p.add_run(text)
        if font_name:
            run.font.name = font_name
        if size:
            run.font.size = Pt(size)
        if bold is not None:
            run.font.bold = bold
        if color:
            run.font.color.rgb = color
        # 中文字体设置
        if font_name:
            r = run._element
            rPr = r.find(qn('w:rPr'))
            if rPr is None:
                rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
                r.insert(0, rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
                rPr.insert(0, rFonts)
            else:
                rFonts.set(qn('w:eastAsia'), font_name)
        return p

    def add_bullet_list(items, font_name="Microsoft YaHei", size=11, color=None):
        for item in items:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(str(item))
            if font_name:
                run.font.name = font_name
            if size:
                run.font.size = Pt(size)
            if color:
                run.font.color.rgb = color

    def add_numbered_list(items, font_name="Microsoft YaHei", size=11, color=None):
        for i, item in enumerate(items, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {item}")
            if font_name:
                run.font.name = font_name
            if size:
                run.font.size = Pt(size)
            if color:
                run.font.color.rgb = color

    def add_table(headers, rows, font_name="Microsoft YaHei", size=10, color=None):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Light Grid Accent 1'
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(h))
            run.font.bold = True
            if font_name:
                run.font.name = font_name
            if size:
                run.font.size = Pt(size)
        for ri, row_data in enumerate(rows):
            for ci, val in enumerate(row_data):
                cell = table.rows[ri + 1].cells[ci]
                cell.text = ""
                run = cell.paragraphs[0].add_run(str(val))
                if font_name:
                    run.font.name = font_name
                if size:
                    run.font.size = Pt(size)
                if color:
                    run.font.color.rgb = color
        doc.add_paragraph()

    # ---- 追加新内容 ----
    default_font = "Microsoft YaHei"
    default_size = 11

    for sec in sections:
        heading = sec.get("heading", "")
        level = sec.get("level", 1)
        content = sec.get("content", "")
        content_type = sec.get("type", "paragraph")

        if heading:
            # 优先使用模板中的标题样式
            heading_style = template_styles.get(f"heading{level}")
            add_styled_paragraph(
                heading,
                style_name=heading_style,
                font_name=default_font,
                size=16 if level == 1 else (14 if level == 2 else 12),
                bold=True,
                space_before=18 if level <= 2 else 12,
                space_after=6
            )

        if content_type == "paragraph":
            if isinstance(content, str):
                for para in content.split("\n"):
                    para = para.strip()
                    if para:
                        add_styled_paragraph(
                            para, font_name=default_font, size=default_size,
                            space_before=2, space_after=4
                        )
            elif isinstance(content, list):
                for para in content:
                    if isinstance(para, str):
                        add_styled_paragraph(
                            para, font_name=default_font, size=default_size,
                            space_before=2, space_after=4
                        )

        elif content_type == "bullet_list":
            if isinstance(content, list):
                add_bullet_list(content, default_font, default_size)

        elif content_type == "numbered_list":
            if isinstance(content, list):
                add_numbered_list(content, default_font, default_size)

        elif content_type == "table":
            headers = sec.get("headers", [])
            rows = sec.get("rows", [])
            if headers and rows:
                add_table(headers, rows, default_font, 10)

    # ---- 保存文件 ----
    if not filename:
        safe_topic = "".join(c if c.isalnum() or c in " _-" else "_" for c in topic)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{safe_topic}_{timestamp}"

    output_path = OUTPUT_DIR / f"{filename}.docx"
    doc.save(str(output_path))
    return {
        "filename": f"{filename}.docx",
        "path": str(output_path),
        "topic": topic,
        "used_template": str(tpl_path),
        "placeholders_replaced": list(template_values.keys()) if template_values else [],
    }


# ============================================================
# 工具 3: 不使用模板，从零生成文档（内置样式）
# ============================================================
@workflow.tool
def generate_docx(
    topic: str,
    sections: list,
    author: str = "",
    style: str = "professional",
    include_toc: bool = True,
    include_page_numbers: bool = True,
    filename: str = "",
    **kwargs
):
    """不使用模板，从零生成带内置样式的 Word 文档。

    Args:
        topic: 文档标题/主题
        sections: 文档章节列表，每个元素为 {"heading": "章节标题", "level": 1-4, "content": "段落文本或列表"}
        author: 作者名称
        style: 文档风格 (professional/creative/simple)
        include_toc: 是否包含目录
        include_page_numbers: 是否包含页码
        filename: 输出文件名（不含扩展名）
    """
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    doc = Document()

    # 样式配置
    style_configs = {
        "professional": {
            "title_font": ("Microsoft YaHei", 22, True, RGBColor(0x1A, 0x1A, 0x2E)),
            "heading1_font": ("Microsoft YaHei", 16, True, RGBColor(0x2C, 0x3E, 0x50)),
            "heading2_font": ("Microsoft YaHei", 13, True, RGBColor(0x34, 0x49, 0x5E)),
            "heading3_font": ("Microsoft YaHei", 12, True, RGBColor(0x2C, 0x3E, 0x50)),
            "body_font": ("Microsoft YaHei", 11, False, RGBColor(0x33, 0x33, 0x33)),
            "accent_color": RGBColor(0x29, 0x80, 0xB9),
            "page_margin": Cm(2.54),
        },
        "creative": {
            "title_font": ("Microsoft YaHei", 26, True, RGBColor(0xE7, 0x4C, 0x3C)),
            "heading1_font": ("Microsoft YaHei", 18, True, RGBColor(0xE7, 0x4C, 0x3C)),
            "heading2_font": ("Microsoft YaHei", 14, True, RGBColor(0xF3, 0x72, 0x59)),
            "heading3_font": ("Microsoft YaHei", 12, True, RGBColor(0xE7, 0x4C, 0x3C)),
            "body_font": ("Microsoft YaHei", 11, False, RGBColor(0x2C, 0x3E, 0x50)),
            "accent_color": RGBColor(0xE7, 0x4C, 0x3C),
            "page_margin": Cm(2.0),
        },
        "simple": {
            "title_font": ("Microsoft YaHei", 20, True, RGBColor(0x33, 0x33, 0x33)),
            "heading1_font": ("Microsoft YaHei", 14, True, RGBColor(0x33, 0x33, 0x33)),
            "heading2_font": ("Microsoft YaHei", 12, True, RGBColor(0x33, 0x33, 0x33)),
            "heading3_font": ("Microsoft YaHei", 11, True, RGBColor(0x33, 0x33, 0x33)),
            "body_font": ("Microsoft YaHei", 10.5, False, RGBColor(0x33, 0x33, 0x33)),
            "accent_color": RGBColor(0x33, 0x33, 0x33),
            "page_margin": Cm(2.54),
        },
    }

    cfg = style_configs.get(style, style_configs["professional"])

    section = doc.sections[0]
    section.top_margin = cfg["page_margin"]
    section.bottom_margin = cfg["page_margin"]
    section.left_margin = cfg["page_margin"]
    section.right_margin = cfg["page_margin"]

    def set_font(run, font_name, size, bold, color):
        run.font.name = font_name
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = color
        r = run._element
        rPr = r.find(qn('w:rPr'))
        if rPr is None:
            rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
            r.insert(0, rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
            rPr.insert(0, rFonts)
        else:
            rFonts.set(qn('w:eastAsia'), font_name)

    def add_formatted_paragraph(text, font_name, size, bold, color,
                                 alignment=None, space_before=0, space_after=6,
                                 first_line_indent=None):
        p = doc.add_paragraph()
        if alignment:
            p.alignment = alignment
        pf = p.paragraph_format
        pf.space_before = Pt(space_before)
        pf.space_after = Pt(space_after)
        if first_line_indent:
            pf.first_line_indent = Cm(first_line_indent)
        run = p.add_run(text)
        set_font(run, font_name, size, bold, color)
        return p

    def add_bullet_list(items, font_name, size, color):
        for item in items:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(str(item))
            set_font(run, font_name, size, False, color)

    def add_numbered_list(items, font_name, size, color):
        for i, item in enumerate(items, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {item}")
            set_font(run, font_name, size, False, color)

    def add_table(headers, rows, font_name, size, color):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Light Grid Accent 1'
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(h))
            set_font(run, font_name, size, True, RGBColor(0xFF, 0xFF, 0xFF))
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2980B9"/>')
            cell._element.get_or_add_tcPr().append(shading)
        for ri, row_data in enumerate(rows):
            for ci, val in enumerate(row_data):
                cell = table.rows[ri + 1].cells[ci]
                cell.text = ""
                run = cell.paragraphs[0].add_run(str(val))
                set_font(run, font_name, size, False, color)
        doc.add_paragraph()

    # 文档标题
    title_font = cfg["title_font"]
    add_formatted_paragraph(
        topic, title_font[0], title_font[1], title_font[2], title_font[3],
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=36, space_after=12
    )

    # 作者和日期
    info_parts = []
    if author:
        info_parts.append(f"作者：{author}")
    info_parts.append(f"日期：{datetime.now().strftime('%Y年%m月%d日')}")
    info_text = "    ".join(info_parts)
    add_formatted_paragraph(
        info_text, cfg["body_font"][0], 10, False, RGBColor(0x7F, 0x8C, 0x8D),
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18
    )
    add_formatted_paragraph(
        "─" * 50, cfg["body_font"][0], 8, False, RGBColor(0xBD, 0xC3, 0xC7),
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12
    )

    # 目录
    if include_toc and sections:
        add_formatted_paragraph(
            "📋 目录", cfg["heading1_font"][0], cfg["heading1_font"][1],
            cfg["heading1_font"][2], cfg["heading1_font"][3],
            space_before=12, space_after=6
        )
        for sec in sections:
            heading = sec.get("heading", "")
            level = sec.get("level", 1)
            indent = "    " * (level - 1)
            add_formatted_paragraph(
                f"{indent}• {heading}",
                cfg["body_font"][0], 10.5, False, RGBColor(0x29, 0x80, 0xB9),
                space_before=2, space_after=2
            )
        doc.add_page_break()

    # 正文章节
    heading_font_map = {
        1: ("heading1_font", cfg["heading1_font"]),
        2: ("heading2_font", cfg["heading2_font"]),
        3: ("heading3_font", cfg["heading3_font"]),
        4: ("body_font", cfg["body_font"]),
    }

    for sec in sections:
        heading = sec.get("heading", "")
        level = sec.get("level", 1)
        content = sec.get("content", "")
        content_type = sec.get("type", "paragraph")

        if heading:
            _, hf = heading_font_map.get(level, heading_font_map[1])
            add_formatted_paragraph(
                heading, hf[0], hf[1], hf[2], hf[3],
                space_before=18 if level <= 2 else 12, space_after=6
            )

        if content_type == "paragraph":
            if isinstance(content, str):
                for para in content.split("\n"):
                    para = para.strip()
                    if para:
                        add_formatted_paragraph(
                            para, cfg["body_font"][0], cfg["body_font"][1],
                            cfg["body_font"][2], cfg["body_font"][3],
                            space_before=2, space_after=4, first_line_indent=0.74
                        )
            elif isinstance(content, list):
                for para in content:
                    if isinstance(para, dict):
                        p_type = para.get("type", "text")
                        p_text = para.get("text", "")
                        if p_type == "text":
                            add_formatted_paragraph(
                                p_text, cfg["body_font"][0], cfg["body_font"][1],
                                cfg["body_font"][2], cfg["body_font"][3],
                                space_before=2, space_after=4
                            )
                        elif p_type == "bold":
                            add_formatted_paragraph(
                                p_text, cfg["body_font"][0], cfg["body_font"][1],
                                True, cfg["accent_color"],
                                space_before=2, space_after=4
                            )
                    else:
                        add_formatted_paragraph(
                            str(para), cfg["body_font"][0], cfg["body_font"][1],
                            cfg["body_font"][2], cfg["body_font"][3],
                            space_before=2, space_after=4
                        )

        elif content_type == "bullet_list":
            if isinstance(content, list):
                add_bullet_list(content, cfg["body_font"][0],
                                cfg["body_font"][1], cfg["body_font"][3])

        elif content_type == "numbered_list":
            if isinstance(content, list):
                add_numbered_list(content, cfg["body_font"][0],
                                  cfg["body_font"][1], cfg["body_font"][3])

        elif content_type == "table":
            headers = sec.get("headers", [])
            rows = sec.get("rows", [])
            if headers and rows:
                add_table(headers, rows, cfg["body_font"][0],
                          cfg["body_font"][1], cfg["body_font"][3])

    # 页码
    if include_page_numbers:
        section = doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._element.append(fldChar1)
        run2 = fp.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._element.append(instrText)
        run3 = fp.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._element.append(fldChar2)

    # 保存
    if not filename:
        safe_topic = "".join(c if c.isalnum() or c in " _-" else "_" for c in topic)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{safe_topic}_{timestamp}"

    output_path = OUTPUT_DIR / f"{filename}.docx"
    doc.save(str(output_path))
    return {"filename": f"{filename}.docx", "path": str(output_path), "topic": topic}


# ============================================================
# LLM 节点 - 文档规划与生成（Agentic 模式）
# ============================================================
SYSTEM_PROMPT = """你是一位专业的 Word 文档生成专家。你的任务是根据用户的需求，生成结构化的 Word 文档。

## 你的核心能力
1. 理解用户需求，规划文档结构和内容
2. 如果用户提供了模板路径，先调用 analyze_template 分析模板
3. 根据模板分析结果，决定使用 generate_from_template 还是 generate_docx
4. 生成专业、清晰的文档章节内容

## 工作流程

### 流程 A: 用户提供了模板 (template_path 不为空)
1. 调用 analyze_template(template_path=...) 分析模板
2. 查看模板中的占位符 {{变量名}}，如果用户提供了 template_values，自动替换
3. 规划需要追加到模板末尾的新内容
4. 调用 generate_from_template 生成文档

### 流程 B: 用户没有提供模板
1. 直接从零规划文档结构
2. 调用 generate_docx 生成文档

## 文档结构规划指南
- 根据用户提供的主题和内容，将文档划分为合理的章节
- 使用多级标题（Heading 1-4）组织层次结构
- 适当使用列表、表格等元素增强可读性
- 确保内容逻辑清晰、语言专业

## sections 参数格式
sections 是一个列表，每个元素是一个字典：

### 段落 (type: "paragraph")
```json
{
    "heading": "章节标题",
    "level": 1,
    "type": "paragraph",
    "content": "段落文本内容..."
}
```

### 无序列表 (type: "bullet_list")
```json
{
    "heading": "要点列表",
    "level": 2,
    "type": "bullet_list",
    "content": ["要点1", "要点2"]
}
```

### 有序列表 (type: "numbered_list")
```json
{
    "heading": "步骤说明",
    "level": 2,
    "type": "numbered_list",
    "content": ["第一步", "第二步"]
}
```

### 表格 (type: "table")
```json
{
    "heading": "数据表格",
    "level": 2,
    "type": "table",
    "headers": ["列1", "列2"],
    "rows": [["数据1", "数据2"]]
}
```

请确保生成的内容专业、准确、有实际价值。
"""

workflow.add_node(LLMNode(
    id="docx_planner",
    agent_style="agentic",
    system_prompt=SYSTEM_PROMPT,
    user_prompt=(
        "请帮我生成一份 Word 文档。\n\n"
        "【文档主题】{topic}\n"
        "【用户提供的详细内容】{content}\n"
        "【模板文件路径】{template_path}\n"
        "【模板占位符替换值】{template_values}\n"
        "【作者】{author}\n"
        "【文档风格（无模板时生效）】{style}\n"
        "【是否包含目录（无模板时生效）】{include_toc}\n"
        "【是否包含页码（无模板时生效）】{include_page_numbers}\n"
        "【输出文件名】{filename}\n\n"
        "请根据是否有模板路径，走对应的工作流程来生成文档。"
    ),
    tools=["analyze_template", "generate_from_template", "generate_docx"],
    stream=True,
    output_to_user=True,
))

workflow.publish()

# ============================================================
# 本地测试入口
# ============================================================
if __name__ == "__main__":
    # 测试无模板模式
    asyncio.run(workflow.run({
        "topic": "项目周报：2024年Q3第10周",
        "content": "",
        "template_path": "",
        "template_values": "",
        "author": "张三",
        "style": "professional",
        "include_toc": "yes",
        "include_page_numbers": "yes",
        "filename": "weekly_report_test",
    }))
